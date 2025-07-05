import streamlit as st
import pandas as pd
import numpy as np
import yfinance as yf
import praw
from pytrends.request import TrendReq
from docx import Document
from datetime import datetime
import openai
import json
import requests
from fastapi import FastAPI, Request
from threading import Thread
from uvicorn import run as uvicorn_run

# ========== CONFIGURATION ==========
REDDIT_CLIENT_ID = "YOUR_REDDIT_CLIENT_ID"
REDDIT_CLIENT_SECRET = "YOUR_REDDIT_CLIENT_SECRET"
REDDIT_USER_AGENT = "delta-ghost-ai-sentiment"
OPENAI_API_KEY = "YOUR_OPENAI_API_KEY"
GEMINI_API_KEY = "YOUR_GEMINI_API_KEY"
GEMINI_URL = "https://generativelanguage.googleapis.com/v1beta/models/gemini-pro:generateContent?key=" + GEMINI_API_KEY

openai.api_key = OPENAI_API_KEY

reddit = praw.Reddit(
    client_id=REDDIT_CLIENT_ID,
    client_secret=REDDIT_CLIENT_SECRET,
    user_agent=REDDIT_USER_AGENT
)

pytrends = TrendReq(hl='en-US', tz=360)

# ========== HELPER FUNCTIONS ==========

def compute_rsi(series, window=14):
    delta = series.diff()
    gain = (delta.where(delta > 0, 0)).rolling(window=window).mean()
    loss = (-delta.where(delta < 0, 0)).rolling(window=window).mean()
    rs = gain / loss
    return 100 - (100 / (1 + rs))

def compute_ema(series, window=21):
    return series.ewm(span=window, adjust=False).mean()

def reddit_sentiment(keyword, limit=50):
    subs = ['stocks', 'options', 'wallstreetbets']
    sentiment = {'positive': 0, 'negative': 0, 'neutral': 0}
    for sub in subs:
        for post in reddit.subreddit(sub).search(keyword, limit=limit):
            title = post.title.lower()
            if "call" in title or "bullish" in title:
                sentiment["positive"] += 1
            elif "put" in title or "bearish" in title:
                sentiment["negative"] += 1
            else:
                sentiment["neutral"] += 1
    return sentiment

def google_trend_score(keyword):
    pytrends.build_payload([keyword], timeframe='now 7-d')
    data = pytrends.interest_over_time()
    if not data.empty:
        return int(data[keyword].iloc[-1])
    return 0

def analyze_stock(ticker):
    df = yf.download(ticker, period='3mo', interval='1d')
    df['RSI'] = compute_rsi(df['Close'])
    df['EMA'] = compute_ema(df['Close'])
    last = df.iloc[-1]
    if last['Close'] > last['EMA'] and last['RSI'] < 70:
        return df, "Buy Call"
    elif last['Close'] < last['EMA'] and last['RSI'] > 30:
        return df, "Buy Put"
    else:
        return df, "Hold"

def generate_openai_summary(ticker, signal, rsi, sentiment, trend):
    prompt = f"""
    Provide a concise, professional summary for a stock trading signal:
    Ticker: {ticker}
    Signal: {signal}
    RSI: {rsi}
    Reddit Sentiment: {sentiment}
    Google Trends Score: {trend}
    Summary should be clear and actionable.
    """
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": prompt}],
        max_tokens=300
    )
    return response.choices[0].message.content.strip()

def generate_gemini_summary(ticker, signal, rsi, sentiment, trend):
    prompt = f"""
    Provide a concise, professional summary for a stock trading signal:
    Ticker: {ticker}
    Signal: {signal}
    RSI: {rsi}
    Reddit Sentiment: {sentiment}
    Google Trends Score: {trend}
    Summary should be clear and actionable.
    """
    payload = {
        "contents": [{"parts": [{"text": prompt}]}]
    }
    headers = {"Content-Type": "application/json"}
    response = requests.post(GEMINI_URL, headers=headers, data=json.dumps(payload))
    return response.json()['candidates'][0]['content']['parts'][0]['text']

# ========== STREAMLIT UI ==========
st.set_page_config(page_title="Delta Ghost AI+Gemini Analyst", layout="centered")
st.title("📈 Delta Ghost Trade Report – Gemini & GPT-4")

ticker = st.text_input("Enter Ticker Symbol", "AAPL")

if st.button("Run Full Analysis"):
    with st.spinner("Analyzing..."):
        df, signal = analyze_stock(ticker)
        rsi = df['RSI'].iloc[-1]
        trend_score = google_trend_score(ticker)
        sentiment = reddit_sentiment(ticker)

        gpt_summary = generate_openai_summary(ticker, signal, rsi, sentiment, trend_score)
        gemini_summary = generate_gemini_summary(ticker, signal, rsi, sentiment, trend_score)

        st.subheader("📊 Indicators")
        st.line_chart(df[['Close', 'EMA']])
        st.write(f"**RSI:** {rsi:.2f} | **Signal:** {signal}")

        st.subheader("📢 Reddit & Trends")
        st.write(sentiment)
        st.write(f"Google Trends Score: {trend_score}")

        st.subheader("🧠 GPT-4 Summary")
        st.info(gpt_summary)

        st.subheader("🧠 Gemini Summary")
        st.success(gemini_summary)

        doc = Document()
        doc.add_heading(f"{ticker} – Trade Report", 0)
        doc.add_paragraph(f"Signal: {signal}\nRSI: {rsi:.2f}")
        doc.add_paragraph(f"Reddit Sentiment: {sentiment}\nGoogle Trends Score: {trend_score}")
        doc.add_paragraph("\nGPT-4 Summary:\n" + gpt_summary)
        doc.add_paragraph("\nGemini Summary:\n" + gemini_summary)
        filename = f"{ticker}_report.docx"
        doc.save(filename)

        with open(filename, "rb") as f:
            st.download_button("📄 Download Report", f, file_name=filename)

# ========== FASTAPI WEBHOOK FOR TRADINGVIEW ==========
app = FastAPI()

@app.post("/webhook")
async def tradingview_webhook(request: Request):
    data = await request.json()
    ticker = data.get("ticker")
    if ticker:
        df, signal = analyze_stock(ticker)
        print(f"Received alert for {ticker}: {signal}")
        # Future: integrate email, Twilio, order execution
    return {"status": "received"}

Thread(target=uvicorn_run, args=(app,), kwargs={"host": "0.0.0.0", "port": 8000}, daemon=True).start()
